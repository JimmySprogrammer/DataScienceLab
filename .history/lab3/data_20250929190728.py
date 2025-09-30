import os
import time
import argparse
import glob
import pandas as pd
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup
from docx import Document
import matplotlib.pyplot as plt

def start_driver(download_dir, attach=False, debugger_address="127.0.0.1:9222"):
    from selenium.webdriver.chrome.options import Options
    options = Options()
    prefs = {"download.default_directory": os.path.abspath(download_dir),
             "download.prompt_for_download": False,
             "profile.default_content_setting_values.automatic_downloads": 1}
    options.add_experimental_option("prefs", prefs)
    options.add_experimental_option("excludeSwitches", ["enable-logging"])
    if attach:
        options.debugger_address = debugger_address
        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
    else:
        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
    return driver

def wait_for_any_download(download_dir, timeout=60):
    end = time.time() + timeout
    while time.time() < end:
        files = glob.glob(os.path.join(download_dir, "*"))
        for f in files:
            if f.endswith(".crdownload") or f.endswith(".part"):
                break
        else:
            if files:
                return max(files, key=os.path.getctime)
        time.sleep(1)
    return None

def click_element_by_text(driver, text, tag="*"):
    xpath = f"//{tag}[contains(normalize-space(.),\"{text}\")]"
    el = WebDriverWait(driver, 15).until(EC.element_to_be_clickable((By.XPATH, xpath)))
    el.click()
    return el

def safe_find_elements_by_xpath(driver, xpath, timeout=8):
    try:
        return WebDriverWait(driver, timeout).until(EC.presence_of_all_elements_located((By.XPATH, xpath)))
    except:
        return []

def get_research_fields_from_ui(driver):
    fields = []
    try:
        add_filter = click_element_by_text(driver, "Filter Results By")
        time.sleep(1)
        popup_options = safe_find_elements_by_xpath(driver, "//div[contains(@class,'popup') or contains(@class,'overlay')]//li|//div[contains(@class,'popup') or contains(@class,'overlay')]//div[@role='option']")
        for el in popup_options:
            txt = el.text.strip()
            if txt:
                fields.append(txt)
    except Exception:
        pass
    if not fields:
        try:
            selects = driver.find_elements(By.TAG_NAME, "select")
            for s in selects:
                try:
                    if "Research" in s.get_attribute("aria-label") or "Research" in s.get_attribute("name"):
                        opts = s.find_elements(By.TAG_NAME, "option")
                        for o in opts:
                            t = o.text.strip()
                            if t:
                                fields.append(t)
                except:
                    continue
        except:
            pass
    fields = list(dict.fromkeys(fields))
    return fields

def apply_results_list_institutions(driver):
    try:
        click_element_by_text(driver, "Results List")
        time.sleep(0.5)
        click_element_by_text(driver, "Institutions")
        time.sleep(1)
        return True
    except:
        return False

def apply_filter_research_field(driver, field_name):
    try:
        click_element_by_text(driver, "Filter Results By")
        time.sleep(0.8)
        candidate = safe_find_elements_by_xpath(driver, f"//*[contains(normalize-space(.),\"{field_name}\")]")
        if candidate:
            candidate[0].click()
            time.sleep(1)
            return True
    except:
        pass
    return False

def export_current_view_csv(driver):
    try:
        export_btns = safe_find_elements_by_xpath(driver, "//button[contains(.,'Export') or contains(@title,'Export') or contains(.,'Download')]")
        if export_btns:
            export_btns[0].click()
            time.sleep(0.5)
            csv_option = safe_find_elements_by_xpath(driver, "//*[contains(text(),'CSV') or contains(text(),'csv') or contains(text(),'XLS')]")
            if csv_option:
                csv_option[0].click()
                return True
    except:
        pass
    return False

def scrape_all_fields(driver, download_dir, max_fields=None):
    saved_files = []
    fields = get_research_fields_from_ui(driver)
    if not fields:
        raise RuntimeError("Cannot detect research fields automatically. Please inspect the page and adjust selectors.")
    if max_fields:
        fields = fields[:max_fields]
    for field in fields:
        applied = apply_filter_research_field(driver, field)
        if not applied:
            print("Failed to apply field:", field)
            continue
        time.sleep(1.2)
        success = export_current_view_csv(driver)
        if not success:
            print("Export click failed for", field)
            continue
        downloaded = wait_for_any_download(download_dir, timeout=60)
        if downloaded:
            base = os.path.basename(downloaded)
            newname = os.path.join(download_dir, f"{field.replace('/', '_')}.csv")
            try:
                os.replace(downloaded, newname)
            except:
                newname = downloaded
            saved_files.append(newname)
            print("Saved:", newname)
        else:
            print("No file downloaded for", field)
    return saved_files

def consolidate_and_analyze(download_dir, institution_names):
    all_files = glob.glob(os.path.join(download_dir, "*.csv"))
    dfs = []
    for f in all_files:
        try:
            df = pd.read_csv(f, dtype=str)
        except:
            try:
                df = pd.read_excel(f, dtype=str)
            except:
                continue
        df["__source_file__"] = os.path.basename(f)
        dfs.append(df)
    if not dfs:
        raise RuntimeError("No CSV/XLS files found in download dir.")
    big = pd.concat(dfs, ignore_index=True, sort=False)
    big.columns = [c.strip() for c in big.columns]
    inst_col = None
    for c in big.columns:
        if c.lower() in ("institution","organizations","organization","institution name"):
            inst_col = c
            break
    if not inst_col:
        possible = [c for c in big.columns if "inst" in c.lower() or "organ" in c.lower()]
        inst_col = possible[0] if possible else big.columns[0]
    results = {}
    for inst in institution_names:
        sub = big[big[inst_col].str.contains(inst, na=False)]
        results[inst] = sub
    return big, results

def make_report(big_df, results_dict, out_prefix="ecnu_report"):
    writer = pd.ExcelWriter(out_prefix + ".xlsx", engine="openpyxl")
    big_df.to_excel(writer, sheet_name="all_data", index=False)
    for name, df in results_dict.items():
        df.to_excel(writer, sheet_name=name[:30], index=False)
    writer.save()
    doc = Document()
    doc.add_heading("ESI Analysis Report", level=1)
    for name, df in results_dict.items():
        doc.add_heading(name, level=2)
        p = doc.add_paragraph(f"Found {len(df)} rows for {name}")
        if len(df) > 0:
            summary = df.iloc[:, :min(6, df.shape[1])]
            tfile = out_prefix + f"_{name[:10]}.png"
            try:
                numeric_rank = None
                for c in df.columns:
                    if "rank" in c.lower():
                        numeric_rank = c
                        break
                if numeric_rank:
                    df2 = df.copy()
                    df2[numeric_rank] = pd.to_numeric(df2[numeric_rank], errors="coerce")
                    df2 = df2.sort_values(numeric_rank).head(20)
                    plt.figure(figsize=(8,4))
                    plt.barh(df2[numeric_rank].astype(str), df2.index)
                    plt.tight_layout()
                    plt.savefig(tfile)
                    plt.close()
                    doc.add_picture(tfile, width=None)
            except Exception:
                pass
    doc.save(out_prefix + ".docx")
    return out_prefix + ".xlsx", out_prefix + ".docx"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--download_dir", default="downloads")
    parser.add_argument("--attach", action="store_true")
    parser.add_argument("--debugger", default="127.0.0.1:9222")
    parser.add_argument("--max_fields", type=int, default=0)
    args = parser.parse_args()
    if not os.path.exists(args.download_dir):
        os.makedirs(args.download_dir, exist_ok=True)
    driver = start_driver(args.download_dir, attach=args.attach, debugger_address=args.debugger)
    print("Please navigate to the ESI Indicators page in the attached Chrome session (if using attach), or let the browser open and navigate manually.")
    input("Press Enter to continue after the ESI Indicators page is visible and you are logged in...")
    applied = apply_results_list_institutions(driver)
    if not applied:
        print("Could not set Results List to Institutions automatically. Please set it manually on the page and press Enter.")
        input("Press Enter after you set Results List -> Institutions...")
    maxf = args.max_fields if args.max_fields>0 else None
    saved = scrape_all_fields(driver, args.download_dir, max_fields=maxf)
    print("Downloaded files:", saved)
    big, results = consolidate_and_analyze(args.download_dir, ["East China Normal University","华东师范大学"])
    out_xlsx, out_docx = make_report(big, results, out_prefix="ecnu_report")
    print("Report files:", out_xlsx, out_docx)
    driver.quit()

if __name__ == "__main__":
    main()
